"""Report builder for Drawing Reviewer — ported from VelocityAI_Platform."""
from __future__ import annotations

import html
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


REPORTS_DIR = Path("data/generated_reports")


@dataclass
class NormalizedFinding:
    id: str
    severity: str
    category: str
    text: str
    recommendation: str = ""
    checked: bool = False


@dataclass
class NormalizedReport:
    drawing_files: list[str]
    overview: dict[str, str]
    summary: str
    findings: list[NormalizedFinding] = field(default_factory=list)
    extracted_data: dict[str, Any] = field(default_factory=dict)
    engineering_notes: list[str] = field(default_factory=list)
    analysis_error: str = ""


def clean_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    if not text:
        return default
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*•]\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\|\s*-+\s*\|", "", text)
    text = text.replace("<br>", " ").replace("</br>", " ")
    return re.sub(r"\n{3,}", "\n\n", text).strip() or default


def _safe(value: Any, default: str = "Not extracted") -> str:
    return clean_text(value, default)


def _coerce_mapping(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, list):
        mapped: dict[str, Any] = {}
        loose_values: list[str] = []
        for index, item in enumerate(value, 1):
            if isinstance(item, dict):
                key = (item.get("key") or item.get("field") or item.get("name") or item.get("label") or item.get("title") or f"item_{index}")
                val = item.get("value") if "value" in item else item.get("text") if "text" in item else item.get("content") if "content" in item else {k: v for k, v in item.items() if k not in {"key", "field", "name", "label", "title"}}
                mapped[str(key)] = val
            elif item not in (None, ""):
                loose_values.append(str(item))
        if loose_values:
            mapped["items"] = loose_values
        return mapped
    if value in (None, ""):
        return {}
    return {"value": value}


def _looks_like_full_markdown_report(text: str) -> bool:
    lowered = text.lower()
    return (
        "| field |" in lowered
        or "|---" in lowered
        or ("drawing overview" in lowered and "key findings" in lowered)
        or lowered.count("##") >= 2
        or len(text) > 2500
    )


def _clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def normalize_findings(check_results: dict[str, Any] | None, checklist: list[dict[str, Any]] | None = None) -> list[NormalizedFinding]:
    raw_items = checklist if checklist is not None else (check_results or {}).get("findings") or []
    findings: list[NormalizedFinding] = []
    for item in raw_items:
        if not isinstance(item, dict):
            continue
        text = clean_text(item.get("text") or item.get("description") or "")
        if not text:
            continue
        findings.append(NormalizedFinding(
            id=str(item.get("id") or f"F{len(findings) + 1}"),
            severity=clean_text(item.get("severity") or "info", "info").lower(),
            category=clean_text(item.get("category") or "General", "General"),
            text=text,
            recommendation=clean_text(item.get("recommendation") or ""),
            checked=bool(item.get("checked", False)),
        ))
    return findings


def build_normalized_report(
    *,
    extracted_data: dict[str, Any],
    check_results: dict[str, Any],
    checklist: list[dict[str, Any]],
    report_sections: dict[str, Any] | None = None,
    drawing_files: list[str] | None = None,
) -> NormalizedReport:
    title_block = _coerce_mapping((extracted_data or {}).get("title_block"))
    overview = {
        "Document Title": _safe(title_block.get("document_title") or title_block.get("title")),
        "Drawing Number": _safe(title_block.get("drawing_number") or title_block.get("drawing_no")),
        "Part Number": _safe(title_block.get("part_number") or title_block.get("part_no")),
        "Revision": _safe(title_block.get("revision") or title_block.get("rev")),
        "Material": _safe(title_block.get("material")),
        "Scale": _safe(title_block.get("scale")),
        "Units": _safe(title_block.get("units")),
        "Review Method": "Automated AI-assisted analysis",
    }
    summary = ""
    if report_sections:
        summary = clean_text(report_sections.get("summary") or report_sections.get("drawing_overview") or "")
    if not summary or _looks_like_full_markdown_report(summary):
        summary = clean_text((check_results or {}).get("summary") or "")
    if not summary:
        summary = "Drawing review completed from extracted drawing data and checklist findings."
    engineering_notes: list[str] = []
    if report_sections:
        raw_notes = report_sections.get("engineering_review_notes") or report_sections.get("recommendations") or []
        if isinstance(raw_notes, str):
            raw_notes = raw_notes.splitlines()
        if isinstance(raw_notes, list):
            engineering_notes = [clean_text(note) for note in raw_notes]
            engineering_notes = [note for note in engineering_notes if note and note not in {"-", "•"}]
    if not engineering_notes:
        engineering_notes = ["AI-generated findings are review assistance only and do not replace engineering approval."]
    return NormalizedReport(
        drawing_files=drawing_files or [],
        overview=overview,
        summary=summary,
        findings=normalize_findings(check_results, checklist),
        extracted_data=extracted_data,
        engineering_notes=engineering_notes,
        analysis_error=clean_text(extracted_data.get("analysis_error") or ""),
    )


def _severity_color(severity: str) -> str:
    return {"critical": "#B00020", "error": "#D84315", "warning": "#B26A00", "info": "#1565C0"}.get((severity or "").lower(), "#374151")


def build_report_preview_html(report: NormalizedReport) -> str:
    reviewed = ", ".join(Path(p).name for p in report.drawing_files) if report.drawing_files else "Not provided"
    overview_rows = "".join(f"<tr><td><b>{html.escape(k)}</b></td><td>{html.escape(v)}</td></tr>" for k, v in report.overview.items())
    if report.findings:
        findings_html = "".join(
            "<tr>"
            f"<td>{html.escape(finding.id)}</td>"
            f"<td><b style='color:{_severity_color(finding.severity)}'>{html.escape(finding.severity)}</b></td>"
            f"<td>{html.escape(finding.category)}</td>"
            f"<td>{html.escape(finding.text)}</td>"
            f"<td>{html.escape(finding.recommendation)}</td>"
            "</tr>"
            for finding in report.findings
        )
    else:
        findings_html = "<tr><td colspan='5'>No checklist findings were generated.</td></tr>"
    notes_html = "".join(f"<li>{html.escape(note)}</li>" for note in report.engineering_notes if note)
    limitation_html = f"<h2>Analysis Limitation</h2><p>{html.escape(report.analysis_error)}</p>" if report.analysis_error else ""
    return f"""
    <div style="font-family: Aptos, Arial, sans-serif; font-size: 13px; line-height: 1.45; padding: 14px;">
      <h1 style="margin: 0 0 6px 0;">Drawing Review Report</h1>
      <p><b>Reviewed drawing file(s):</b> {html.escape(reviewed)}</p>
      <h2>Drawing Overview</h2>
      <table border="1" cellspacing="0" cellpadding="6" style="border-collapse: collapse; width: 100%;"><tbody>{overview_rows}</tbody></table>
      <h2>Review Summary</h2>
      <p>{html.escape(report.summary)}</p>
      <h2>Key Findings</h2>
      <table border="1" cellspacing="0" cellpadding="6" style="border-collapse: collapse; width: 100%;">
        <thead><tr style="background: #1F4E78; color: white;"><th>ID</th><th>Severity</th><th>Category</th><th>Finding</th><th>Recommendation</th></tr></thead>
        <tbody>{findings_html}</tbody>
      </table>
      <h2>Engineering Review Notes</h2>
      <ul>{notes_html}</ul>
      {limitation_html}
      <p><b>Human engineering review required before release or manufacturing use.</b></p>
    </div>
    """


def _set_default_styles(doc: Document) -> None:
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10)
    except Exception:
        pass


def _add_docx_overview(doc: Document, report: NormalizedReport) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Details"
    for f, v in report.overview.items():
        row = table.add_row().cells
        row[0].text = f
        row[1].text = v


def _add_docx_findings(doc: Document, report: NormalizedReport) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, header in enumerate(["ID", "Severity", "Category", "Finding", "Recommendation"]):
        table.rows[0].cells[idx].text = header
    if not report.findings:
        row = table.add_row().cells
        row[0].text = "-"; row[1].text = "info"; row[2].text = "General"
        row[3].text = "No checklist findings were generated."
        row[4].text = "Perform manual engineering review before release."
        return
    for finding in report.findings:
        row = table.add_row().cells
        row[0].text = finding.id
        row[1].text = finding.severity
        row[2].text = finding.category
        row[3].text = finding.text
        row[4].text = finding.recommendation


def _format_structured_item(item: Any) -> str:
    if isinstance(item, dict):
        clean = {str(k): clean_text(v) for k, v in item.items() if clean_text(v)}
        return "; ".join(f"{k}: {v}" for k, v in clean.items()) if clean else ""
    return clean_text(item)


def _add_docx_list_section(doc: Document, title: str, values: Any) -> None:
    doc.add_heading(title, level=2)
    if not values:
        doc.add_paragraph("No data extracted.")
        return
    added = False
    if isinstance(values, dict):
        for key, value in values.items():
            text = _format_structured_item(value)
            if text:
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {text}", style="List Bullet")
                added = True
    elif isinstance(values, list):
        for item in values[:100]:
            text = _format_structured_item(item)
            if text:
                doc.add_paragraph(text, style="List Bullet")
                added = True
    else:
        text = clean_text(values)
        if text:
            doc.add_paragraph(text)
            added = True
    if not added:
        doc.add_paragraph("No data extracted.")


def build_drawing_review_docx(
    *,
    extracted_data: dict[str, Any],
    check_results: dict[str, Any],
    checklist: list[dict[str, Any]],
    report_sections: dict[str, Any] | None = None,
    drawing_files: list[str] | None = None,
    output_dir: str | Path | None = None,
) -> str:
    report = build_normalized_report(
        extracted_data=extracted_data, check_results=check_results,
        checklist=checklist, report_sections=report_sections, drawing_files=drawing_files,
    )
    output_root = Path(output_dir) if output_dir else REPORTS_DIR
    output_root.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_default_styles(doc)
    title = doc.add_heading("Drawing Review Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("AI-assisted engineering drawing review")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if report.drawing_files:
        doc.add_paragraph("Reviewed drawing file(s): " + ", ".join(Path(p).name for p in report.drawing_files))
    doc.add_heading("1. Drawing Overview", level=1)
    _add_docx_overview(doc, report)
    doc.add_heading("2. Review Summary", level=1)
    doc.add_paragraph(report.summary)
    doc.add_heading("3. Key Findings", level=1)
    _add_docx_findings(doc, report)
    doc.add_heading("4. Extracted Drawing Data", level=1)
    _add_docx_list_section(doc, "Title Block", extracted_data.get("title_block"))
    _add_docx_list_section(doc, "Revision Block", extracted_data.get("revision_block"))
    _add_docx_list_section(doc, "Views", extracted_data.get("views"))
    _add_docx_list_section(doc, "Dimensions", extracted_data.get("dimensions"))
    _add_docx_list_section(doc, "Tolerances", extracted_data.get("tolerances"))
    _add_docx_list_section(doc, "GD&T", extracted_data.get("gdt"))
    _add_docx_list_section(doc, "Annotations", extracted_data.get("annotations"))
    _add_docx_list_section(doc, "General Notes", extracted_data.get("general_notes"))
    doc.add_heading("5. Engineering Review Notes", level=1)
    for note in report.engineering_notes:
        doc.add_paragraph(note, style="List Bullet")
    if report.analysis_error:
        doc.add_heading("6. Analysis Limitation", level=1)
        doc.add_paragraph(report.analysis_error)
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("Human engineering review required before release or manufacturing use.").bold = True
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    first_drawing = Path(report.drawing_files[0]).stem if report.drawing_files else "drawing"
    safe_stem = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in first_drawing)[:80]
    output_path = output_root / f"{safe_stem}_review_report_{timestamp}.docx"
    doc.save(str(output_path))
    return str(output_path)


def build_drawing_review_pdf(
    *,
    extracted_data: dict[str, Any],
    check_results: dict[str, Any],
    checklist: list[dict[str, Any]],
    report_sections: dict[str, Any] | None = None,
    drawing_files: list[str] | None = None,
    output_dir: str | Path | None = None,
) -> str:
    """Generate PDF report. Falls back to empty string if reportlab is not installed."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except ImportError:
        return ""

    report = build_normalized_report(
        extracted_data=extracted_data, check_results=check_results,
        checklist=checklist, report_sections=report_sections, drawing_files=drawing_files,
    )
    output_root = Path(output_dir) if output_dir else REPORTS_DIR
    output_root.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    first_drawing = Path(report.drawing_files[0]).stem if report.drawing_files else "drawing"
    safe_stem = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in first_drawing)[:80]
    output_path = output_root / f"{safe_stem}_review_report_{timestamp}.pdf"

    styles = getSampleStyleSheet()
    story = [
        Paragraph("Drawing Review Report", styles["Title"]),
        Paragraph("AI-assisted engineering drawing review", styles["Normal"]),
        Spacer(1, 12),
    ]
    overview_rows = [["Field", "Details"]] + [[k, v] for k, v in report.overview.items()]
    tbl = Table(overview_rows, colWidths=[120, 350])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.extend([Paragraph("Drawing Overview", styles["Heading1"]), tbl, Spacer(1, 12)])
    story.extend([Paragraph("Review Summary", styles["Heading1"]), Paragraph(html.escape(report.summary), styles["BodyText"]), Spacer(1, 12)])
    story.append(Paragraph("Key Findings", styles["Heading1"]))
    for finding in report.findings:
        story.append(Paragraph(f"<b>[{html.escape(finding.severity)}] {html.escape(finding.category)}</b> — {html.escape(finding.text)}", styles["BodyText"]))
        if finding.recommendation:
            story.append(Paragraph(f"Recommendation: {html.escape(finding.recommendation)}", styles["Italic"]))
        story.append(Spacer(1, 6))
    story.append(Paragraph("Human engineering review required before release or manufacturing use.", styles["Heading2"]))
    SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36).build(story)
    return str(output_path)